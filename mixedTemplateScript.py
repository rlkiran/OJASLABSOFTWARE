from __future__ import print_function
from mailmerge import MailMerge
# from datetime import date
from docx import Document
# from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

template = "Template002.docx"
filename = 'test.docx'

TestList = []

patientData = {
    'PatientName': 'PatientName',
    'PatientAge': 'PatientAge',
    'PatientSex': 'PatientSex',
    'PatientId': 'PatientId',
    'PatientPhone': 'PatientPhone',
    'SpecimenID': 'SpecimenID',
    'DateCollected': 'DateCollected',
    'DateReceived': 'DateReceived',
    'DateofReport': 'DateofReport',
    'SampleType': 'SampleType',
    'Doctor': 'Doctor',
    'Phone': 'Phone'}

DocsDirectory = os.getcwd() + "\\Docs"


def setStyle(documentData):
    obj_styles = documentData.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'


def addTestTables(docLoc, filename, *Ttwo):
    TestList = list(Ttwo)
    TestName = TestList[0]['TestName']
    os.chdir(docLoc)
    listOfFiles = os.listdir(os.getcwd())
    if filename not in listOfFiles:
        document = Document()
        document.save(filename)
    document = Document(filename)
    try:
        setStyle(document)
    except Exception as e:
        print(str(e))
    table = document.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    hdr_cellMain = table.rows[0].cells
    hdrTitle = hdr_cellMain[0].add_paragraph()
    hdrTitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdrTitle.add_run(TestName, style='CommentsStyle').bold = True
    table.cell(0, 0).merge(table.cell(0, 4))
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = 'PARAMETER'
    hdr_cells[1].text = 'TECHNOLOGY'
    hdr_cells[2].text = 'RESULT'
    hdr_cells[3].text = 'UNITS'
    hdr_cells[4].text = 'REFERENCE RANGE'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[2].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[3].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[4].paragraphs[0].runs[0]
    run.font.bold = True
    for col in Ttwo:
        row_cells = table.add_row().cells
        row_cells[0].text = str(col.get('Parameter'))
        row_cells[1].text = str(col.get('Technology'))
        row_cells[2].text = str(col.get('result'))
        row_cells[3].text = str(col.get('Units'))
        row_cells[4].text = str(col.get('RefRange'))
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)
        document.save(filename)


def addPatientDetails(template, docLoc, filename, **pd):
    documentOne = MailMerge(template)
    documentOne.merge(
        PatientName=pd['PatientName'],
        PatientAge=pd['PatientAge'],
        PatientSex=pd['PatientSex'],
        PatientId=pd['PatientId'],
        PatientPhone=pd['PatientPhone'],
        SpecimenID=pd['SpecimenID'],
        DateCollected=pd['DateCollected'],
        DateReceived=pd['DateReceived'],
        DateofReport=pd['DateofReport'],
        SampleType=pd['SampleType'],
        Doctor=pd['Doctor'],
        Phone=pd['Phone'])
    os.chdir(docLoc)
    documentOne.write(filename)


def addEoR(docloc, filename):
    os.chdir(docloc)
    document = Document(filename)
    try:
        obj_styles = document.styles
        obj_charstyle = obj_styles.add_style('endStyle', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(12)
        obj_font.name = 'Times New Roman'
    except Exception as e:
        print(str(e))
    endText = document.add_paragraph()
    # endText.alignment = WD_ALIGN_PARAGRAPH.CENTER
    endText.add_run('\n**END OF REPORT**', style='endStyle').bold = True
    document.save(filename)

# addPatientDetails(**patientData.copy())
# addTestTables(*TestList.copy())
# addTestTables(*TestList.copy())
# addEoR()
# os.startfile(filename)
